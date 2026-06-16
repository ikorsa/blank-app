"""Markdown to DOCX conversion for medical reports."""

from __future__ import annotations

import re
from io import BytesIO

from docx import Document
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def _clean_text(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = text.replace("—", "-")
    text = text.replace("–", "-")
    return text.strip()


def _strip_inline_markup(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return _clean_text(text)


def _set_run_font(run, *, bold: bool = False, italic: bool = False, mono: bool = False, size: int = 12) -> None:
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if mono:
        run.font.name = "Courier New"
    else:
        run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _format_body_paragraph(paragraph) -> None:
    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(6)


def _add_title(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run(_strip_inline_markup(text))
    _set_run_font(run, bold=True, size=16)


def _add_rich_paragraph(document: Document, text: str, *, style: str | None = None, justify: bool = True) -> None:
    paragraph = document.add_paragraph(style=style)
    if justify and style is None:
        _format_body_paragraph(paragraph)
    text = _clean_text(text)
    parts = re.split(r"(\*\*.+?\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            _set_run_font(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            _set_run_font(run, mono=True)
        else:
            run = paragraph.add_run(part)
            _set_run_font(run)


def _is_table_row(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and "|" in stripped[1:]


def _is_separator_row(line: str) -> bool:
    return bool(re.fullmatch(r"\|?[\s:\-|]+\|?", line.strip()))


def _parse_table_row(line: str) -> list[str]:
    return [_strip_inline_markup(cell.strip()) for cell in line.strip().strip("|").split("|")]


def _set_table_column_widths(table, col_count: int) -> None:
    total = Cm(16.0)
    if col_count == 4:
        weights = [2.2, 2.2, 2.8, 3.8]
    elif col_count == 3:
        weights = [2.5, 3.5, 4.0]
    elif col_count == 2:
        weights = [3.5, 6.5]
    else:
        weights = [1.0] * col_count
    weight_sum = sum(weights)
    for idx, column in enumerate(table.columns):
        column.width = int(total * weights[idx] / weight_sum)


def build_markdown_docx(markdown_text: str, title: str) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    for heading_level in range(1, 4):
        style = document.styles[f"Heading {heading_level}"]
        style.font.name = "Times New Roman"

    _add_title(document, title)

    lines = markdown_text.splitlines()
    i = 0
    skip_first_h1 = True

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped or stripped == "---":
            i += 1
            continue

        if stripped.startswith("# ") and not stripped.startswith("## "):
            heading = _strip_inline_markup(stripped[2:])
            if skip_first_h1 and heading == title:
                skip_first_h1 = False
                i += 1
                continue
            document.add_heading(heading, level=1)
            i += 1
            continue

        if stripped.startswith("## "):
            document.add_heading(_strip_inline_markup(stripped[3:]), level=2)
            i += 1
            continue

        if stripped.startswith("### "):
            document.add_heading(_strip_inline_markup(stripped[4:]), level=3)
            i += 1
            continue

        if _is_table_row(stripped):
            table_rows: list[list[str]] = []
            while i < len(lines) and _is_table_row(lines[i].strip()):
                row_line = lines[i].strip()
                if not _is_separator_row(row_line):
                    table_rows.append(_parse_table_row(row_line))
                i += 1
            if table_rows:
                table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                table.style = "Table Grid"
                table.autofit = False
                _set_table_column_widths(table, len(table_rows[0]))
                for r_idx, row in enumerate(table_rows):
                    for c_idx, cell_text in enumerate(row):
                        cell = table.rows[r_idx].cells[c_idx]
                        cell.text = ""
                        paragraph = cell.paragraphs[0]
                        _format_body_paragraph(paragraph)
                        run = paragraph.add_run(cell_text)
                        _set_run_font(run, bold=(r_idx == 0), size=11)
                document.add_paragraph()
            continue

        if stripped.startswith("- "):
            _add_rich_paragraph(document, stripped[2:], style="List Bullet", justify=False)
            i += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered:
            _add_rich_paragraph(document, numbered.group(2), style="List Number", justify=False)
            i += 1
            continue

        if stripped.startswith("> "):
            quote = stripped[2:]
            if "python scripts" in quote.lower() or "docs/MEDICAL" in quote:
                i += 1
                continue
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(1)
            paragraph.paragraph_format.right_indent = Cm(0.5)
            _format_body_paragraph(paragraph)
            run = paragraph.add_run(_strip_inline_markup(quote))
            _set_run_font(run, italic=True, size=11)
            i += 1
            continue

        _add_rich_paragraph(document, stripped)
        i += 1

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
